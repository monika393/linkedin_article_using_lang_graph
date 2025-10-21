"""
Export utilities for LinkedIn Article Generation System
Handles Word document and JPEG image export functionality
"""

import os
import re
from datetime import datetime
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFont
import requests
from io import BytesIO
import openai
from openai import OpenAI


def export_to_word(article_data: Dict[str, Any], output_path: str) -> str:
    """
    Export article to Word document format
    
    Args:
        article_data: Dictionary containing article data
        output_path: Path where to save the Word document
        
    Returns:
        Path to the saved Word document
    """
    doc = Document()
    
    # Set document title
    title = doc.add_heading(article_data['topic'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_heading('Article Information', level=1)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('Revisions Made:', str(article_data['revisions_made'])),
        ('Word Count:', str(len(article_data['article'].split()))),
        ('Character Count:', str(len(article_data['article'])))
    ]
    
    for i, (key, value) in enumerate(info_data):
        info_table.cell(i, 0).text = key
        info_table.cell(i, 1).text = value
    
    # Add main article content
    doc.add_heading('Article Content', level=1)
    
    # Parse and format the article content
    article_text = article_data.get('article', '')
    lines = article_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if it's a heading (starts with ##)
        if line.startswith('##'):
            heading_text = line.replace('##', '').strip()
            doc.add_heading(heading_text, level=2)
        elif line.startswith('#'):
            heading_text = line.replace('#', '').strip()
            doc.add_heading(heading_text, level=1)
        else:
            # Regular paragraph
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add LinkedIn post section
    doc.add_heading('LinkedIn Post', level=1)
    doc.add_paragraph(article_data['linkedin_post'])
    
    # Add hashtags
    doc.add_heading('Hashtags', level=1)
    hashtags_text = ', '.join(article_data['hashtags'])
    doc.add_paragraph(hashtags_text)
    
    # Add SEO keywords
    doc.add_heading('SEO Keywords', level=1)
    keywords_text = ', '.join(article_data['seo_keywords'])
    doc.add_paragraph(keywords_text)
    
    # Add image prompt
    doc.add_heading('Image Generation Prompt', level=1)
    doc.add_paragraph(article_data['image_prompt'])
    
    # Save the document
    doc.save(output_path)
    return output_path


def create_placeholder_image(text: str, width: int = 1200, height: int = 630) -> Image.Image:
    """
    Create a high-quality placeholder image with professional design
    
    Args:
        text: Text to display on the image
        width: Image width in pixels
        height: Image height in pixels
        
    Returns:
        PIL Image object
    """
    # Create image with professional gradient background
    image = Image.new('RGB', (width, height), color='#0f172a')
    draw = ImageDraw.Draw(image)
    
    # Create a sophisticated gradient background based on content
    if "green" in text.lower() or "sustainability" in text.lower():
        # Green theme for sustainability
        for y in range(height):
            green_intensity = int(20 + (y / height) * 60)
            blue_intensity = int(30 + (y / height) * 40)
            color = (green_intensity, green_intensity + 30, blue_intensity)
            draw.line([(0, y), (width, y)], fill=color)
    elif "optimization" in text.lower() or "efficiency" in text.lower():
        # Blue theme for optimization
        for y in range(height):
            blue_intensity = int(20 + (y / height) * 50)
            purple_intensity = int(30 + (y / height) * 40)
            color = (blue_intensity, blue_intensity + 20, blue_intensity + 40)
            draw.line([(0, y), (width, y)], fill=color)
    else:
        # Default blue-to-purple gradient
        for y in range(height):
            blue_intensity = int(15 + (y / height) * 40)
            purple_intensity = int(20 + (y / height) * 30)
            color = (blue_intensity, blue_intensity + 20, blue_intensity + 40)
            draw.line([(0, y), (width, y)], fill=color)
    
    # Add subtle geometric patterns and visual elements
    for i in range(0, width, 120):
        for j in range(0, height, 120):
            if (i + j) % 240 == 0:
                # Add subtle circles
                draw.ellipse([i, j, i+15, j+15], fill=(255, 255, 255, 20), outline=None)
            elif (i + j) % 360 == 0:
                # Add subtle squares
                draw.rectangle([i+5, j+5, i+10, j+10], fill=(255, 255, 255, 15), outline=None)
    
    # Add some connecting lines for visual interest
    for i in range(0, width, 200):
        for j in range(0, height, 150):
            if (i + j) % 300 == 0:
                draw.line([i, j, i+50, j+30], fill=(255, 255, 255, 10), width=1)
    
    # Add text with better typography
    try:
        # Try to use a system font
        title_font = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 56)
        subtitle_font = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 32)
    except:
        try:
            title_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 56)
            subtitle_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 32)
        except:
            title_font = ImageFont.load_default()
            subtitle_font = ImageFont.load_default()
    
    # Clean and format the text for better display
    clean_text = text.replace("Create a high-definition abstract image that", "").replace("encapsulates", "").strip()
    clean_text = clean_text.replace("Create an HD abstract image", "").replace("that symbolizes", "").strip()
    
    # Extract key concepts for visual representation
    if "sustainability" in clean_text.lower() or "carbon" in clean_text.lower():
        title_text = "Green AI"
        subtitle_text = "Sustainable Technology"
    elif "optimization" in clean_text.lower() or "efficiency" in clean_text.lower():
        title_text = "AI Optimization"
        subtitle_text = "Performance & Efficiency"
    elif "machine learning" in clean_text.lower() or "llm" in clean_text.lower():
        title_text = "AI Innovation"
        subtitle_text = "Machine Learning"
    else:
        # Use first few words as title
        words = clean_text.split()
        title_words = words[:6] if len(words) > 6 else words
        subtitle_words = words[6:] if len(words) > 6 else []
        
        title_text = " ".join(title_words)
        subtitle_text = " ".join(subtitle_words) if subtitle_words else "Technology"
    
    # Draw title
    title_bbox = draw.textbbox((0, 0), title_text, font=title_font)
    title_width = title_bbox[2] - title_bbox[0]
    title_x = (width - title_width) // 2
    title_y = height // 2 - 60
    
    # Add title shadow
    draw.text((title_x + 3, title_y + 3), title_text, font=title_font, fill='#000000')
    # Add main title
    draw.text((title_x, title_y), title_text, font=title_font, fill='#ffffff')
    
    # Draw subtitle
    if subtitle_text:
        subtitle_bbox = draw.textbbox((0, 0), subtitle_text, font=subtitle_font)
        subtitle_width = subtitle_bbox[2] - subtitle_bbox[0]
        subtitle_x = (width - subtitle_width) // 2
        subtitle_y = title_y + 80
        
        # Add subtitle shadow
        draw.text((subtitle_x + 2, subtitle_y + 2), subtitle_text, font=subtitle_font, fill='#000000')
        # Add main subtitle
        draw.text((subtitle_x, subtitle_y), subtitle_text, font=subtitle_font, fill='#e2e8f0')
    
    # Add a subtle border
    draw.rectangle([10, 10, width-10, height-10], outline='#3b82f6', width=2)
    
    return image


def generate_dalle_image(prompt: str, output_path: str) -> str:
    """
    Generate image using DALL-E API
    
    Args:
        prompt: Text prompt for image generation
        output_path: Path where to save the generated image
        
    Returns:
        Path to the saved image
    """
    try:
        # Initialize OpenAI client
        client = OpenAI()
        
        # Generate image using DALL-E
        response = client.images.generate(
            model="dall-e-3",
            prompt=prompt,
            size="1792x1024",  # LinkedIn-optimized size
            quality="hd",
            n=1
        )
        
        # Get image URL
        image_url = response.data[0].url
        
        # Download and save the image
        image_response = requests.get(image_url)
        image_response.raise_for_status()
        
        # Save as JPEG
        with open(output_path, 'wb') as f:
            f.write(image_response.content)
        
        return output_path
        
    except Exception as e:
        print(f"⚠️  DALL-E generation failed: {e}")
        print("   Falling back to placeholder image...")
        
        # Fallback to placeholder
        return create_placeholder_fallback(prompt, output_path)


def create_placeholder_fallback(prompt: str, output_path: str) -> str:
    """
    Create a high-quality placeholder image when DALL-E fails
    
    Args:
        prompt: Text prompt for image generation
        output_path: Path where to save the image
        
    Returns:
        Path to the saved image
    """
    image = create_placeholder_image(prompt)
    image.save(output_path, 'JPEG', quality=95, optimize=True)
    return output_path


def export_image_to_jpeg(article_data: Dict[str, Any], output_path: str, 
                        use_dalle: bool = True) -> str:
    """
    Export image as JPEG file
    
    Args:
        article_data: Dictionary containing article data
        output_path: Path where to save the JPEG image
        use_placeholder: If True, create a placeholder image; if False, try to download from URL
        
    Returns:
        Path to the saved JPEG image
    """
    if use_dalle:
        # Use DALL-E to generate the image
        image_prompt = article_data.get('image_prompt', 'Professional LinkedIn article image')
        
        print(f"🎨 Generating DALL-E image with prompt: {image_prompt[:100]}...")
        
        try:
            return generate_dalle_image(image_prompt, output_path)
        except Exception as e:
            print(f"❌ DALL-E generation failed: {e}")
            print("   Using high-quality placeholder instead...")
            return create_placeholder_fallback(image_prompt, output_path)
    else:
        # Use placeholder image
        image_prompt = article_data.get('image_prompt', 'LinkedIn Article Image')
        return create_placeholder_fallback(image_prompt, output_path)
    
    return output_path


def create_placeholder_fallback(image_prompt: str, output_path: str) -> str:
    """
    Create a high-quality placeholder image when DALL-E fails
    
    Args:
        image_prompt: The prompt that was used for image generation
        output_path: Path where to save the placeholder image
        
    Returns:
        Path to the saved placeholder image
    """
    # Create a professional placeholder image
    image = create_placeholder_image(image_prompt)
    
    # Save as JPEG
    image.save(output_path, 'JPEG', quality=95)
    print(f"✅ Placeholder image saved to: {output_path}")
    
    return output_path


def create_article_package(article_data: Dict[str, Any], output_dir: str) -> Dict[str, str]:
    """
    Create a complete article package with Word document and JPEG image
    
    Args:
        article_data: Dictionary containing article data
        output_dir: Directory where to save the files
        
    Returns:
        Dictionary with paths to created files
    """
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate safe filename from topic
    safe_topic = re.sub(r'[^\w\s-]', '', article_data['topic']).strip()
    safe_topic = re.sub(r'[-\s]+', '_', safe_topic)[:50]
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    # Create file paths
    word_path = os.path.join(output_dir, f"{safe_topic}_{timestamp}.docx")
    image_path = os.path.join(output_dir, f"{safe_topic}_{timestamp}.jpg")
    
    # Export files
    word_file = export_to_word(article_data, word_path)
    image_file = export_image_to_jpeg(article_data, image_path, use_dalle=True)
    
    return {
        'word_document': word_file,
        'image_file': image_file,
        'output_directory': output_dir
    }
